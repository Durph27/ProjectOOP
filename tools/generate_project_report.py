from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
ASSET_DIR = OUT_DIR / "assets"
OUT_FILE = OUT_DIR / "BAO_CAO_DU_AN_OOP.docx"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


def font(size=30, bold=False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path("C:/Windows/Fonts") / name), size)


def multiline_center(draw, box, text, fnt, fill="#17202A", spacing=7):
    x1, y1, x2, y2 = box
    max_chars = max(8, int((x2 - x1) / (fnt.size * 0.56)))
    lines = []
    for paragraph in text.split("\n"):
        lines.extend(wrap(paragraph, max_chars) or [""])
    bbox = draw.multiline_textbbox((0, 0), "\n".join(lines), font=fnt, spacing=spacing, align="center")
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - w) / 2, (y1 + y2 - h) / 2),
        "\n".join(lines),
        font=fnt,
        fill=fill,
        spacing=spacing,
        align="center",
    )


def box(draw, coords, text, fill="#EAF3F8", outline="#1F4E78", title=False):
    draw.rounded_rectangle(coords, radius=18, fill=fill, outline=outline, width=4)
    multiline_center(draw, coords, text, font(30, title))


def arrow(draw, start, end, color="#365F91", width=5):
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        p = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line([end, p], fill=color, width=width)


def canvas(title):
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 55), title, font=font(48, True), fill="#17365D", anchor="ma")
    return image, draw


def make_diagrams():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    image, draw = canvas("KIẾN TRÚC TỔNG THỂ HỆ THỐNG")
    box(draw, (650, 145, 1150, 255), "Người dùng / JavaFX MainFrame", title=True)
    box(draw, (70, 390, 500, 520), "CollectionService\nDataCollectorFactory", fill="#E2F0D9")
    box(draw, (685, 390, 1115, 520), "PreprocessingService\nPreprocessorChain", fill="#FFF2CC")
    box(draw, (1300, 390, 1730, 520), "AnalysisService\nAnalyzerRegistry", fill="#FCE4D6")
    box(draw, (70, 720, 500, 880), "CSV Collector\nSocialMediaPost\n(Các collector mở rộng)")
    box(draw, (685, 720, 1115, 880), "HTML → URL → Emoji\n→ Chuẩn hóa → Stopword")
    box(draw, (1300, 690, 1730, 910), "SentimentModelFactory\nPython API / Dictionary\n4 Analyzer nghiệp vụ")
    for x in (285, 900, 1515):
        arrow(draw, (900, 255), (x, 390))
    arrow(draw, (285, 520), (285, 720))
    arrow(draw, (900, 520), (900, 720))
    arrow(draw, (1515, 520), (1515, 690))
    image.save(ASSET_DIR / "architecture.png")

    image, draw = canvas("BIỂU ĐỒ PACKAGE VÀ QUAN HỆ PHỤ THUỘC")
    coords = {
        "ui": (690, 140, 1110, 250),
        "service": (690, 330, 1110, 440),
        "collector": (70, 570, 450, 690),
        "preprocessor": (500, 570, 880, 690),
        "sentiment": (930, 570, 1310, 690),
        "analyzer": (1360, 570, 1740, 690),
        "config": (300, 870, 680, 990),
        "model": (1080, 870, 1460, 990),
    }
    colors = ["#D9EAF7", "#E2F0D9", "#FFF2CC", "#FCE4D6"]
    for i, (name, c) in enumerate(coords.items()):
        box(draw, c, f"com.humanitarian.{name}", fill=colors[i % len(colors)], title=True)
    arrow(draw, (900, 250), (900, 330))
    for name in ("collector", "preprocessor", "sentiment", "analyzer"):
        c = coords[name]
        arrow(draw, (900, 440), ((c[0] + c[2]) // 2, c[1]))
    for name in ("collector", "preprocessor", "sentiment", "analyzer", "service", "ui"):
        c = coords[name]
        arrow(draw, ((c[0] + c[2]) // 2, c[3]), (1270, 870), color="#7F8C8D", width=3)
    arrow(draw, (490, 870), (750, 690), color="#7F8C8D", width=3)
    arrow(draw, (490, 870), (1120, 690), color="#7F8C8D", width=3)
    image.save(ASSET_DIR / "packages.png")

    image, draw = canvas("BIỂU ĐỒ LỚP CỐT LÕI")
    class_boxes = {
        "Analyzer<T>\n<<interface>>\n+ analyze(List<Post>): T": (60, 150, 470, 330),
        "AnalyzerRegistry\n- Map<String, Analyzer<?>>\n+ register(), get()": (680, 150, 1120, 350),
        "AnalysisService\n- registry\n- sentimentFactory\n+ assignSentiment()\n+ runAnalyzer<T>()": (1320, 150, 1760, 390),
        "SentimentModelProvider\n<<interface>>\n+ analyze(): SentimentResult": (60, 650, 500, 830),
        "SentimentModelFactory\n- Map<String, Provider>\n+ getConfigured()": (680, 620, 1120, 820),
        "SocialMediaPost\n- rawContent, content\n- sentiment, confidence": (1320, 650, 1760, 830),
    }
    for text, c in class_boxes.items():
        box(draw, c, text, fill="#EAF3F8", title=False)
    arrow(draw, (470, 240), (680, 240))
    arrow(draw, (1120, 250), (1320, 250))
    arrow(draw, (500, 740), (680, 720))
    arrow(draw, (1120, 720), (1320, 740))
    arrow(draw, (1540, 390), (1540, 650))
    draw.text((560, 205), "quản lý", font=font(24), fill="#365F91")
    draw.text((1170, 210), "được dùng bởi", font=font(24), fill="#365F91")
    image.save(ASSET_DIR / "classes.png")

    image, draw = canvas("LUỒNG XÁC ĐỊNH SENTIMENT CỤC BỘ")
    items = [
        ("Văn bản đã tiền xử lý", "#D9EAF7"),
        ("SentimentModelFactory.getConfigured()", "#E2F0D9"),
        ("Python API khả dụng?", "#FFF2CC"),
        ("PhoBERT tải thành công?\nTokenizer → logits → softmax → nhãn", "#FCE4D6"),
        ("Fallback rule-based\nTừ điển + phủ định + tăng cường", "#EDE7F6"),
        ("SentimentResult\nnhãn + confidence + modelName", "#D9EAD3"),
    ]
    ys = [130, 300, 470, 640, 850, 1030]
    for (text, color), y in zip(items, ys):
        h = 125 if y != 640 else 150
        box(draw, (480, y, 1320, y + h), text, fill=color, title=y in (130, 1030))
    for a, b in zip(ys[:-1], ys[1:]):
        h = 150 if a == 640 else 125
        arrow(draw, (900, a + h), (900, b))
    draw.text((1335, 515), "Không → Java dictionary", font=font(25), fill="#943634")
    draw.text((1335, 700), "Không → Python rule-based", font=font(25), fill="#943634")
    image.save(ASSET_DIR / "sentiment_flow.png")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], BLUE)
        set_cell_text(table.rows[0].cells[i], h, True, WHITE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if len(table.rows) % 2 == 1:
                shade(cells[i], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    doc.add_paragraph()
    return table


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Mục lục sẽ được cập nhật khi mở Word (Ctrl+A, F9)."
    separate.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, end])


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_code(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.style = "Code"
        p.add_run(line)


def add_figure(doc, path, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.style = "Caption"


def page_break(doc):
    doc.add_page_break()


def placeholder_page(doc, heading, prompt):
    add_heading(doc, heading, 1)
    p = doc.add_paragraph(prompt)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(7)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(127, 127, 127)
    page_break(doc)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BLUE)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("365F91")
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)
    styles["Caption"].font.italic = True
    styles["Caption"].font.size = Pt(10)
    code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Cm(0.7)
    code_style.paragraph_format.space_after = Pt(0)

    header = section.header.paragraphs[0]
    header.text = "BÁO CÁO DỰ ÁN LẬP TRÌNH HƯỚNG ĐỐI TƯỢNG"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].font.name = "Times New Roman"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    add_page_number(section.footer.paragraphs[0])

    # Trang giới thiệu để trống
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(8)
    r = p.add_run("TRANG GIỚI THIỆU")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Times New Roman"
    p2 = doc.add_paragraph("(Nhóm tự bổ sung thông tin trường, học phần, đề tài, giảng viên và thành viên)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].italic = True
    p2.runs[0].font.color.rgb = RGBColor(127, 127, 127)
    page_break(doc)

    add_heading(doc, "MỤC LỤC", 1)
    add_toc(doc)
    page_break(doc)

    placeholder_page(doc, "LỜI NÓI ĐẦU", "(Nhóm tự viết lời nói đầu tại đây)")
    placeholder_page(doc, "PHÂN CÔNG CÔNG VIỆC", "(Nhóm tự bổ sung bảng phân công công việc thành viên tại đây)")

    add_heading(doc, "1. TỔNG QUAN VỀ DỰ ÁN", 1)
    add_heading(doc, "1.1. Bối cảnh và mục tiêu", 2)
    add_body(doc, "Dự án xây dựng ứng dụng desktop phân tích dữ liệu mạng xã hội phục vụ logistics nhân đạo trong bối cảnh thiên tai, với cấu hình hiện tại tập trung vào bão Yagi tại miền Bắc Việt Nam. Hệ thống hỗ trợ nhập dữ liệu, tiền xử lý văn bản tiếng Việt, xác định sentiment và tổng hợp kết quả thành các biểu đồ phục vụ quan sát, đánh giá.")
    add_bullets(doc, [
        "Theo dõi biến động tâm lý công chúng theo thời gian.",
        "Phân loại các loại thiệt hại được đề cập trong bài đăng.",
        "Đánh giá mức hài lòng theo từng loại hàng cứu trợ.",
        "Theo dõi sentiment của từng loại cứu trợ theo thời gian.",
    ])
    add_heading(doc, "1.2. Phạm vi triển khai hiện tại", 2)
    add_body(doc, "Nguồn dữ liệu hoạt động thực tế là tệp CSV. Các collector Twitter/X, Facebook, TikTok và YouTube đã có khung lớp để mở rộng nhưng chưa kết nối API thật. AI được dùng cho bài toán sentiment; phân loại thiệt hại và loại cứu trợ hiện dùng đối sánh từ khóa cấu hình trong JSON.")
    add_heading(doc, "1.3. Công nghệ sử dụng", 2)
    add_table(doc, ["Công nghệ", "Vai trò"], [
        ("Java 17", "Ngôn ngữ chính, triển khai mô hình OOP và nghiệp vụ."),
        ("JavaFX 21", "Giao diện desktop, bảng dữ liệu và biểu đồ."),
        ("Maven", "Quản lý dependency, build và đóng gói ứng dụng."),
        ("Gson / OpenCSV / OkHttp", "Đọc JSON, nhập CSV và gọi REST API Python."),
        ("SLF4J + Logback", "Ghi log và theo dõi lỗi."),
        ("Python + Flask", "Cung cấp REST API sentiment chạy cục bộ."),
        ("Transformers + PyTorch + PhoBERT", "Suy luận sentiment tiếng Việt bằng mô hình học máy."),
        ("JSON configuration", "Tách danh mục, từ khóa và lựa chọn provider khỏi mã nguồn."),
    ], [4, 12])
    page_break(doc)

    placeholder_page(doc, "2. THỐNG KÊ DỮ LIỆU THU THẬP", "(Để trống cho đến khi nhóm chốt bộ dữ liệu; nên bổ sung số bài đăng, thời gian, nguồn, tỷ lệ sentiment và phân bố danh mục)")

    add_heading(doc, "3. KIẾN TRÚC VÀ BIỂU ĐỒ UML", 1)
    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    add_figure(doc, ASSET_DIR / "architecture.png", "Hình 3.1. Kiến trúc tổng thể của hệ thống")
    add_body(doc, "Kiến trúc chia hệ thống thành các tầng trách nhiệm rõ ràng. UI gọi các service; service điều phối factory, registry và các strategy; model là định dạng dữ liệu chung xuyên suốt. Python API được cách ly sau SentimentModelProvider nên phần còn lại của Java không phụ thuộc trực tiếp vào Flask hoặc PhoBERT.")
    page_break(doc)

    add_heading(doc, "3.2. Biểu đồ package", 2)
    add_figure(doc, ASSET_DIR / "packages.png", "Hình 3.2. Các package và quan hệ phụ thuộc chính")
    add_table(doc, ["Package", "Chức năng, nhiệm vụ"], [
        ("com.humanitarian", "Điểm vào chương trình và vòng đời JavaFX."),
        ("ui", "Nhận thao tác người dùng, hiển thị bảng, biểu đồ và nhận xét."),
        ("service", "Điều phối các use case thu thập, tiền xử lý và phân tích."),
        ("collector", "Chuẩn hóa việc thu thập dữ liệu từ nhiều nguồn về SocialMediaPost."),
        ("preprocessor", "Các bước làm sạch văn bản và chuỗi xử lý có thứ tự."),
        ("sentiment", "Abstraction, factory và implementation xác định sentiment."),
        ("analyzer", "Bốn bài toán tổng hợp dữ liệu, dùng generic interface Analyzer<T>."),
        ("config", "Đọc JSON và cung cấp cấu hình dùng chung."),
        ("model", "Các entity/DTO đại diện bài đăng và kết quả phân tích."),
    ], [4, 12])
    page_break(doc)

    add_heading(doc, "3.3. Biểu đồ lớp cốt lõi", 2)
    add_figure(doc, ASSET_DIR / "classes.png", "Hình 3.3. Quan hệ giữa các lớp điều phối, registry, provider và model")
    add_body(doc, "AnalysisService là lớp điều phối trung tâm của giai đoạn phân tích. Lớp lấy SentimentModelProvider từ factory để gán sentiment, sau đó lấy Analyzer<T> từ registry để tạo kết quả có kiểu phù hợp. SocialMediaPost lưu cả nội dung gốc, nội dung đã xử lý và kết quả sentiment để dùng lại giữa nhiều analyzer.")
    add_heading(doc, "3.4. Các interface và lớp trừu tượng quan trọng", 2)
    add_table(doc, ["Thành phần", "Kiểu", "Ý nghĩa thiết kế"], [
        ("DataCollector", "Interface", "Hợp đồng chung cho các nguồn dữ liệu."),
        ("AbstractCollector", "Abstract class", "Cố định quy trình collect và dành doCollect cho lớp con."),
        ("TextPreprocessor", "Interface", "Mỗi bước tiền xử lý là một strategy độc lập."),
        ("SentimentModelProvider", "Interface", "Cho phép thay provider Python API hoặc từ điển Java."),
        ("Analyzer<T>", "Generic interface", "Mỗi analyzer trả một kiểu kết quả riêng qua cùng hợp đồng."),
    ], [4, 3, 9])
    page_break(doc)

    add_heading(doc, "4. THIẾT KẾ LỚP VÀ NHIỆM VỤ", 1)
    add_heading(doc, "4.1. Nhóm lớp collector", 2)
    add_table(doc, ["Lớp", "Nhiệm vụ"], [
        ("DataCollectorFactory", "Đăng ký và trả collector theo khóa nền tảng; kết hợp Factory và Registry."),
        ("AbstractCollector", "Kiểm tra đầu vào, ghi log và định nghĩa Template Method collect → doCollect."),
        ("CsvFileCollector", "Đọc CSV UTF-8, parse nhiều định dạng ngày, lọc thời gian và tạo SocialMediaPost."),
        ("Twitter/Facebook/Tiktok/YoutubeCollector", "Khung mở rộng cho các nguồn mạng xã hội; chưa gọi API thật."),
    ], [5, 11])
    add_heading(doc, "4.2. Nhóm lớp tiền xử lý", 2)
    add_table(doc, ["Lớp", "Nhiệm vụ"], [
        ("PreprocessorChain", "Giữ List<TextPreprocessor>, sắp xếp theo order, bật/tắt và xử lý lỗi từng bước."),
        ("HtmlCleanerPreprocessor", "Loại bỏ HTML tag và entity."),
        ("UrlRemover", "Loại bỏ URL bằng biểu thức chính quy."),
        ("EmojiHandler", "Chuyển emoji phổ biến thành từ khóa cảm xúc."),
        ("VietnameseNormalizer", "Chuẩn hóa Unicode, chữ thường, khoảng trắng và ký tự lặp."),
        ("StopWordRemover", "Loại các từ ít giá trị phân tích bằng HashSet."),
    ], [5, 11])
    add_heading(doc, "4.3. Nhóm lớp sentiment", 2)
    add_table(doc, ["Lớp", "Nhiệm vụ"], [
        ("SentimentModelFactory", "Registry provider và chọn provider theo config; fallback sang dictionary khi Python API không khả dụng."),
        ("PythonApiSentimentProvider", "Adapter chuyển lời gọi Java thành HTTP JSON tới Flask API."),
        ("DictionaryBasedSentimentProvider", "Thuật toán sentiment cục bộ bằng từ điển Java."),
        ("SentimentResult", "DTO chứa nhãn, confidence và tên mô hình."),
    ], [5, 11])
    add_heading(doc, "4.4. Nhóm lớp analyzer", 2)
    add_table(doc, ["Analyzer", "Kiểu kết quả", "Nhiệm vụ"], [
        ("SentimentTimelineAnalyzer", "Map<LocalDate, TimeSentimentData>", "Đếm positive/negative/neutral theo ngày."),
        ("DamageClassificationAnalyzer", "Map<CategoryDefinition, List<DamageReport>>", "Phân loại đa nhãn theo từ khóa thiệt hại."),
        ("ReliefSatisfactionAnalyzer", "Map<CategoryDefinition, ReliefSentiment>", "Tính tỷ lệ hài lòng theo loại cứu trợ."),
        ("ReliefSentimentTimelineAnalyzer", "Map<CategoryDefinition, Map<LocalDate, TimeSentimentData>>", "Theo dõi sentiment từng loại cứu trợ theo ngày."),
    ], [4.5, 6, 6])
    page_break(doc)

    add_heading(doc, "5. CÁC KỸ THUẬT LẬP TRÌNH ĐÃ ÁP DỤNG", 1)
    add_table(doc, ["Kỹ thuật", "Áp dụng ở đâu", "Lợi ích"], [
        ("Factory + Registry", "DataCollectorFactory, SentimentModelFactory, AnalyzerRegistry", "Tập trung đăng ký/chọn implementation; thêm hoặc thay thế thành phần mà ít sửa service/UI."),
        ("Singleton instance", "AppConfig và ba factory/registry", "Đảm bảo toàn ứng dụng dùng chung một cấu hình và một tập đăng ký."),
        ("Generic interface", "Analyzer<T>, runAnalyzer<T>()", "Dùng chung hợp đồng nhưng giữ kiểu kết quả riêng cho từng bài toán."),
        ("Strategy", "DataCollector, TextPreprocessor, SentimentModelProvider, Analyzer<T>", "Tách thuật toán khỏi nơi sử dụng và hỗ trợ thay thế."),
        ("Template Method", "AbstractCollector.collect()/doCollect()", "Tái sử dụng kiểm tra và logging; lớp con chỉ cài đặt phần đặc thù."),
        ("Pipeline", "PreprocessorChain", "Thêm, bỏ, sắp xếp và bật/tắt từng bước làm sạch."),
        ("Adapter", "PythonApiSentimentProvider", "Che giấu HTTP/JSON và giữ contract Java ổn định."),
        ("Configuration-driven", "AppConfig + các tệp JSON", "Thay danh mục, từ khóa, provider và bước xử lý không cần sửa code."),
        ("Exception handling", "Tải model Python, đọc CSV, gọi API, chạy preprocessor/analyzer", "Cô lập lỗi và duy trì khả năng hoạt động bằng fallback."),
    ], [4, 6, 6])
    add_heading(doc, "5.1. Java Collection Framework trong analyzer và hạ tầng", 2)
    add_table(doc, ["Cấu trúc", "Vị trí tiêu biểu", "Lý do chọn"], [
        ("List", "Danh sách bài đăng, từ khóa, processor, DamageReport", "Giữ tập phần tử có thứ tự và cho phép lặp tuần tự."),
        ("Map", "Factory/registry và kết quả analyzer", "Ánh xạ khóa → implementation hoặc khóa → kết quả."),
        ("LinkedHashMap", "AnalyzerRegistry và kết quả danh mục", "Giữ thứ tự đăng ký/cấu hình để hiển thị ổn định."),
        ("HashMap", "DataCollectorFactory, SentimentModelFactory, EmojiHandler", "Tra cứu trung bình O(1), không yêu cầu thứ tự."),
        ("TreeMap", "Timeline theo LocalDate", "Tự động sắp xếp ngày tăng dần, thuận tiện vẽ biểu đồ."),
        ("Set / HashSet", "Từ điển sentiment và stopword", "Kiểm tra membership nhanh, loại trùng lặp."),
    ], [3, 7, 6])
    add_heading(doc, "5.2. Try-catch và cơ chế fallback", 2)
    add_body(doc, "Khi Python khởi tạo SentimentModel, khối try-catch thử tải tokenizer và PhoBERT. Nếu dependency, mạng hoặc cache gặp lỗi, hệ thống đặt use_ml = false và chuyển sang rule-based. Ở Java, SentimentModelFactory kiểm tra isAvailable() của PythonApiSentimentProvider; nếu API không chạy, factory trả DictionaryBasedSentimentProvider. PreprocessorChain và AnalysisService cũng bắt lỗi theo từng bước để một lỗi cục bộ không làm dừng toàn bộ pipeline.")
    page_break(doc)

    add_heading(doc, "6. BỐN THUỘC TÍNH CỦA LẬP TRÌNH HƯỚNG ĐỐI TƯỢNG", 1)
    add_body(doc, "Cách trình bày dưới đây dùng thuật ngữ nhất quán với lecture note IT3100: trừu tượng hóa (abstraction), đóng gói (encapsulation), kế thừa (inheritance) và đa hình (polymorphism).")
    add_heading(doc, "6.1. Trừu tượng hóa - Abstraction", 2)
    add_body(doc, "Các interface mô tả khả năng mà không buộc caller biết chi tiết thực hiện: DataCollector mô tả thu thập, TextPreprocessor mô tả xử lý văn bản, SentimentModelProvider mô tả xác định sentiment và Analyzer<T> mô tả bài toán phân tích. Ví dụ, AnalysisService chỉ gọi provider.analyze(text), không cần biết kết quả đến từ từ điển Java hay PhoBERT qua Flask.")
    add_heading(doc, "6.2. Đóng gói - Encapsulation", 2)
    add_body(doc, "Các model giữ thuộc tính ở mức private và cung cấp getter/setter. Quan trọng hơn, TimeSentimentData đóng gói hành vi incrementPositive(), incrementNegative(), incrementNeutral(); ReliefSentiment tự tính lại satisfactionRate khi bộ đếm thay đổi. Nhờ đó quy tắc nhất quán nằm trong chính đối tượng thay vì bị lặp ở UI hoặc analyzer.")
    add_heading(doc, "6.3. Kế thừa - Inheritance", 2)
    add_body(doc, "Các collector cụ thể kế thừa AbstractCollector. Lớp cha cung cấp quy trình collect chung, còn lớp con override doCollect để triển khai nguồn dữ liệu cụ thể. HumanitarianApplication extends Application và MainFrame extends BorderPane để tham gia vòng đời, hệ thống layout của JavaFX.")
    add_heading(doc, "6.4. Đa hình - Polymorphism", 2)
    add_body(doc, "Một biến kiểu DataCollector có thể tham chiếu CsvFileCollector hoặc collector mạng xã hội; một biến SentimentModelProvider có thể tham chiếu PythonApiSentimentProvider hoặc DictionaryBasedSentimentProvider. Khi gọi cùng phương thức analyze() hoặc collect(), implementation thực tế quyết định hành vi. Đây là đa hình thông qua interface và overriding.")
    add_heading(doc, "6.5. Liên hệ SOLID", 2)
    add_bullets(doc, [
        "SRP: mỗi collector, preprocessor, provider và analyzer có một trách nhiệm chính.",
        "OCP: mở rộng bằng implementation và đăng ký mới thay vì sửa toàn bộ luồng.",
        "LSP: các implementation có thể thay thế abstraction tương ứng.",
        "ISP: các interface nhỏ, tập trung theo một nhóm hành vi.",
        "DIP: service làm việc chủ yếu qua abstraction, dù việc lấy Singleton trực tiếp vẫn làm giảm khả năng test.",
    ])
    page_break(doc)

    add_heading(doc, "7. THUẬT TOÁN XÁC ĐỊNH SENTIMENT JAVA LOCALLY", 1)
    add_figure(doc, ASSET_DIR / "sentiment_flow.png", "Hình 7.1. Luồng lựa chọn mô hình và fallback sentiment")
    add_heading(doc, "7.1. Dữ liệu và cấu trúc", 2)
    add_body(doc, "DictionaryBasedSentimentProvider dùng bốn HashSet: positiveWords, negativeWords, negationWords và intensifierWords. HashSet phù hợp vì thuật toán cần kiểm tra một từ/cụm từ có thuộc từ điển hay không với chi phí trung bình O(1). Dữ liệu đầu vào là content đã qua PreprocessorChain.")
    add_heading(doc, "7.2. Các bước thuật toán", 2)
    add_bullets(doc, [
        "Chuyển văn bản về chữ thường và tách theo khoảng trắng.",
        "Duyệt từng từ; đồng thời tạo bigram gồm từ hiện tại và từ kế tiếp.",
        "Nếu gặp từ phủ định, bật cờ negated; nếu gặp từ tăng cường, đặt hệ số intensifier = 1.5.",
        "Từ/cụm tích cực cộng positiveScore; khi có phủ định thì chuyển điểm sang negativeScore.",
        "Từ/cụm tiêu cực cộng negativeScore; khi có phủ định thì cộng 0.5 vào positiveScore.",
        "Duyệt thêm các cụm từ nhiều tiếng xuất hiện trong toàn văn bản và cộng 1.5 điểm.",
        "So sánh hai tổng điểm để trả POSITIVE, NEGATIVE hoặc NEUTRAL.",
    ])
    add_code(doc, """
totalScore = positiveScore + negativeScore
if totalScore == 0:
    sentiment = NEUTRAL; confidence = 0.5
elif positiveScore > negativeScore:
    sentiment = POSITIVE
    confidence = min(0.95, 0.5 + (positiveScore-negativeScore)/(totalScore*2))
elif negativeScore > positiveScore:
    sentiment = NEGATIVE
    confidence = min(0.95, 0.5 + (negativeScore-positiveScore)/(totalScore*2))
else:
    sentiment = NEUTRAL; confidence = 0.5
""")
    add_heading(doc, "7.3. Ví dụ minh họa", 2)
    add_table(doc, ["Văn bản", "Phân tích", "Kết quả dự kiến"], [
        ("Cứu trợ đến rất nhanh và đầy đủ", "“rất” tăng cường; “nhanh”, “đầy đủ” tích cực", "POSITIVE"),
        ("Hàng cứu trợ chưa đến, người dân rất lo lắng", "“chưa đến”, “lo lắng” tiêu cực; “rất” tăng cường", "NEGATIVE"),
        ("Hôm nay trời nhiều mây", "Không khớp từ điển", "NEUTRAL, confidence 0.5"),
    ], [6, 6, 4])
    add_heading(doc, "7.4. Đánh giá", 2)
    add_body(doc, "Ưu điểm của thuật toán là chạy hoàn toàn offline, nhanh, dễ giải thích và luôn sẵn sàng làm fallback. Hạn chế là không hiểu đầy đủ ngữ cảnh, mỉa mai, quan hệ phủ định xa hoặc từ mới ngoài từ điển. Confidence là điểm quy ước, không phải xác suất đã hiệu chỉnh như softmax của PhoBERT.")
    page_break(doc)

    add_heading(doc, "8. LUỒNG XỬ LÝ CÁC BÀI TOÁN PHÂN TÍCH", 1)
    add_heading(doc, "8.1. Luồng nghiệp vụ chung", 2)
    add_code(doc, """
CSV → List<SocialMediaPost>
    → PreprocessingService.preprocess(posts)
    → AnalysisService.assignSentiment(posts)
    → AnalyzerRegistry.get(analyzerId)
    → analyzer.analyze(posts)
    → JavaFX chart + nhận xét
""")
    add_heading(doc, "8.2. Phân loại thiệt hại", 2)
    add_body(doc, "DamageClassificationAnalyzer nạp các CategoryDefinition từ damage-categories.json, duyệt từng bài và đếm keyword khớp bằng content.contains(keyword). Một bài có thể thuộc nhiều danh mục. Confidence được tính theo công thức min(0.95, 0.3 + matchCount × 0.15).")
    add_heading(doc, "8.3. Mức hài lòng theo loại cứu trợ", 2)
    add_body(doc, "ReliefSatisfactionAnalyzer xác định loại cứu trợ bằng từ khóa, sau đó dùng sentiment đã gán để tăng bộ đếm. Tỷ lệ hài lòng bằng positiveCount / totalCount. Việc tách sentiment khỏi analyzer giúp thay mô hình sentiment mà không viết lại thuật toán tổng hợp.")
    add_heading(doc, "8.4. Tổng hợp theo thời gian", 2)
    add_body(doc, "SentimentTimelineAnalyzer và ReliefSentimentTimelineAnalyzer sử dụng TreeMap<LocalDate, ...>. computeIfAbsent tạo bộ đếm khi gặp ngày mới; TreeMap giữ ngày có thứ tự tăng dần nên dữ liệu sẵn sàng cho biểu đồ.")
    page_break(doc)

    add_heading(doc, "9. GIỚI HẠN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_table(doc, ["Giới hạn hiện tại", "Hướng phát triển"], [
        ("Chỉ CSV hoạt động thực tế.", "Kết nối API thật, thêm retry, pagination và rate limit."),
        ("Damage/relief dùng keyword matching.", "Tách CategoryClassifier và bổ sung mô hình multi-label."),
        ("AnalysisService gọi sentiment từng bài.", "Sử dụng true batching theo batchSize."),
        ("Singleton làm tăng global state.", "Áp dụng constructor injection để dễ unit test."),
        ("MainFrame chứa nhiều trách nhiệm.", "Tách view, controller/view-model và chart builder."),
        ("Analyzer.isEnabled() mặc định luôn true.", "Liên kết trực tiếp trạng thái enabled với AppConfig."),
    ], [8, 8])
    add_heading(doc, "10. KẾT LUẬN", 1)
    add_body(doc, "Dự án thể hiện rõ tư duy lập trình hướng đối tượng qua việc nhận diện các điểm có thể thay đổi và đặt chúng sau interface. Factory/registry, Singleton, generic interface, Strategy, Template Method, Pipeline và Adapter phối hợp để tạo hệ thống có khả năng mở rộng, thay thế mô hình sentiment và bổ sung bài toán phân tích. Thiết kế hiện tại phù hợp phạm vi đồ án, đồng thời vẫn chỉ ra được các hướng cải tiến khi hệ thống phát triển lớn hơn.")
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    add_bullets(doc, [
        "SOICT, Lecture Note IT3100 - Object-Oriented Programming, các phần về abstraction, encapsulation, inheritance, polymorphism, interface, generic, exception handling và Java Collection Framework.",
        "Mã nguồn dự án ProjectOOP tại thời điểm tạo báo cáo.",
        "Tài liệu chính thức Java 17, JavaFX, Flask, PyTorch và Hugging Face Transformers.",
    ])

    # Repeat headers for every table and improve pagination behavior.
    for table in doc.tables:
        set_repeat_table_header(table.rows[0])
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    make_diagrams()
    print(build_document())
